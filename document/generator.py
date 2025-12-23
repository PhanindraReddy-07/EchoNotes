"""
Smart Document Generator
========================
Generates structured documents from transcripts with:
- Executive Summary
- Key Sentences
- Key Concepts
- Study Questions
- Related Topics
- Full Content

Supports: Markdown, PDF, DOCX, HTML
"""
import re
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime


@dataclass
class KeyConcept:
    """A key concept extracted from content"""
    term: str
    definition: str
    frequency: int = 1
    importance: float = 0.5
    
    def to_dict(self) -> Dict:
        return {
            'term': self.term,
            'definition': self.definition,
            'frequency': self.frequency,
            'importance': round(self.importance, 2)
        }


@dataclass
class StudyQuestion:
    """A study/review question"""
    question: str
    question_type: str  # 'factual', 'conceptual', 'application', 'analysis'
    difficulty: str     # 'easy', 'medium', 'hard'
    suggested_answer: Optional[str] = None
    
    def to_dict(self) -> Dict:
        return {
            'question': self.question,
            'type': self.question_type,
            'difficulty': self.difficulty,
            'answer': self.suggested_answer
        }


@dataclass
class DocumentContent:
    """Structured document content"""
    title: str
    executive_summary: str
    key_sentences: List[str]
    key_concepts: List[KeyConcept]
    study_questions: List[StudyQuestion]
    related_topics: List[str]
    full_content: str
    
    # Meeting-specific (optional)
    action_items: List[str] = field(default_factory=list)
    decisions: List[str] = field(default_factory=list)
    deadlines: List[str] = field(default_factory=list)
    
    # Metadata
    source: str = ""
    duration: float = 0.0
    word_count: int = 0
    generated_at: str = ""
    
    def to_dict(self) -> Dict:
        return {
            'title': self.title,
            'executive_summary': self.executive_summary,
            'key_sentences': self.key_sentences,
            'key_concepts': [c.to_dict() for c in self.key_concepts],
            'study_questions': [q.to_dict() for q in self.study_questions],
            'related_topics': self.related_topics,
            'action_items': self.action_items,
            'decisions': self.decisions,
            'deadlines': self.deadlines,
            'metadata': {
                'source': self.source,
                'duration': self.duration,
                'word_count': self.word_count,
                'generated_at': self.generated_at
            }
        }


class ContentAnalyzer:
    """
    Analyzes transcript content to extract structured information
    """
    
    # Topic indicators for different domains
    TOPIC_KEYWORDS = {
        'technology': ['software', 'app', 'code', 'api', 'database', 'server', 'algorithm', 'data', 'system', 'platform', 'digital', 'computer', 'network', 'programming'],
        'business': ['revenue', 'profit', 'market', 'customer', 'sales', 'strategy', 'growth', 'investment', 'budget', 'roi', 'kpi', 'stakeholder'],
        'science': ['research', 'study', 'experiment', 'hypothesis', 'theory', 'analysis', 'method', 'result', 'evidence', 'scientific'],
        'education': ['learn', 'student', 'course', 'lecture', 'exam', 'curriculum', 'knowledge', 'skill', 'training', 'education'],
        'social_media': ['instagram', 'facebook', 'twitter', 'tiktok', 'youtube', 'post', 'share', 'followers', 'viral', 'content', 'hashtag', 'like'],
        'health': ['health', 'medical', 'treatment', 'patient', 'disease', 'symptom', 'diagnosis', 'therapy', 'medicine', 'doctor'],
    }
    
    # Question templates by type
    QUESTION_TEMPLATES = {
        'factual': [
            "What is {concept}?",
            "Define {concept}.",
            "What are the main features of {concept}?",
            "Who/What is responsible for {concept}?",
            "When was {concept} introduced/mentioned?",
        ],
        'conceptual': [
            "Why is {concept} important?",
            "How does {concept} work?",
            "What is the relationship between {concept} and {related}?",
            "Explain the significance of {concept}.",
        ],
        'application': [
            "How can {concept} be applied in real-world scenarios?",
            "Give an example of {concept} in practice.",
            "How would you use {concept} to solve a problem?",
        ],
        'analysis': [
            "Compare and contrast {concept} with {related}.",
            "What are the advantages and disadvantages of {concept}?",
            "Analyze the impact of {concept}.",
            "What conclusions can be drawn about {concept}?",
        ],
    }
    
    def __init__(self):
        self.stopwords = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
            'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
            'could', 'should', 'may', 'might', 'must', 'shall', 'can', 'need',
            'this', 'that', 'these', 'those', 'it', 'its', 'they', 'them', 'their',
            'we', 'us', 'our', 'you', 'your', 'he', 'him', 'his', 'she', 'her',
            'i', 'me', 'my', 'as', 'if', 'then', 'so', 'than', 'such', 'when',
            'where', 'which', 'who', 'what', 'how', 'all', 'each', 'every',
            'both', 'few', 'more', 'most', 'other', 'some', 'any', 'no', 'not',
            'only', 'own', 'same', 'just', 'also', 'very', 'even', 'back', 'now',
        }
    
    def analyze(self, text: str, title: str = "Document") -> DocumentContent:
        """
        Analyze text and extract structured content
        
        Args:
            text: Input transcript/text
            title: Document title
            
        Returns:
            DocumentContent with all sections
        """
        # Clean text
        clean_text = self._clean_text(text)
        sentences = self._split_sentences(clean_text)
        
        # Extract components
        executive_summary = self._generate_summary(sentences)
        key_sentences = self._extract_key_sentences(sentences)
        key_concepts = self._extract_concepts(clean_text)
        study_questions = self._generate_questions(key_concepts, sentences)
        related_topics = self._find_related_topics(clean_text)
        
        return DocumentContent(
            title=title,
            executive_summary=executive_summary,
            key_sentences=key_sentences,
            key_concepts=key_concepts,
            study_questions=study_questions,
            related_topics=related_topics,
            full_content=clean_text,
            word_count=len(clean_text.split()),
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove headers and metadata from transcript files
        lines = text.split('\n')
        content_lines = []
        in_content = False
        
        for line in lines:
            line = line.strip()
            # Skip metadata lines
            if any(x in line.lower() for x in ['==', '--', 'transcript:', 'audio:', 'duration:', 'confidence:', 'words:', 'timestamps:']):
                if 'transcript:' in line.lower():
                    in_content = True
                continue
            if line.startswith('[') and line.endswith(']'):
                # Remove timestamp markers but keep content
                line = re.sub(r'\[\d{2}:\d{2}\]', '', line).strip()
            if line:
                content_lines.append(line)
        
        text = ' '.join(content_lines)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if len(s.strip()) > 10]
    
    def _generate_summary(self, sentences: List[str], max_sentences: int = 3) -> str:
        """Generate executive summary using extractive method"""
        if not sentences:
            return ""
        
        if len(sentences) <= max_sentences:
            return ' '.join(sentences)
        
        # Score sentences by importance
        scored = []
        for i, sent in enumerate(sentences):
            score = self._score_sentence(sent, i, len(sentences))
            scored.append((score, sent))
        
        # Get top sentences
        scored.sort(reverse=True)
        top_sentences = [s for _, s in scored[:max_sentences]]
        
        # Reorder by original position
        ordered = []
        for sent in sentences:
            if sent in top_sentences:
                ordered.append(sent)
                if len(ordered) == max_sentences:
                    break
        
        return ' '.join(ordered)
    
    def _score_sentence(self, sentence: str, position: int, total: int) -> float:
        """Score sentence importance"""
        score = 0.0
        
        # Position score (first and last sentences more important)
        if position == 0:
            score += 0.3
        elif position == total - 1:
            score += 0.1
        elif position < total * 0.2:
            score += 0.2
        
        # Length score (medium length preferred)
        words = sentence.split()
        if 10 <= len(words) <= 30:
            score += 0.2
        
        # Keyword score
        important_words = ['important', 'key', 'main', 'significant', 'essential', 
                         'primary', 'major', 'allows', 'enables', 'provides']
        for word in important_words:
            if word in sentence.lower():
                score += 0.1
        
        return score
    
    def _extract_key_sentences(self, sentences: List[str], num: int = 5) -> List[str]:
        """Extract most important sentences"""
        if len(sentences) <= num:
            return sentences
        
        scored = []
        for i, sent in enumerate(sentences):
            score = self._score_sentence(sent, i, len(sentences))
            scored.append((score, sent))
        
        scored.sort(reverse=True)
        return [s for _, s in scored[:num]]
    
    def _extract_concepts(self, text: str, max_concepts: int = 10) -> List[KeyConcept]:
        """Extract key concepts/terms from text"""
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Count word frequencies (excluding stopwords)
        freq = {}
        for word in words:
            if word not in self.stopwords:
                freq[word] = freq.get(word, 0) + 1
        
        # Find noun phrases (simple approach)
        noun_phrases = re.findall(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b', text)
        for phrase in noun_phrases:
            phrase_lower = phrase.lower()
            if phrase_lower not in self.stopwords:
                freq[phrase_lower] = freq.get(phrase_lower, 0) + 2  # Boost phrases
        
        # Also find compound terms
        compound_patterns = [
            r'\b(\w+\s+service)\b',
            r'\b(\w+\s+platform)\b',
            r'\b(\w+\s+system)\b',
            r'\b(\w+\s+network)\b',
            r'\b(social\s+\w+)\b',
        ]
        for pattern in compound_patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                freq[match] = freq.get(match, 0) + 3  # Boost compounds
        
        # Sort by frequency and get top concepts
        sorted_terms = sorted(freq.items(), key=lambda x: -x[1])
        
        concepts = []
        seen = set()
        
        for term, count in sorted_terms:
            if len(concepts) >= max_concepts:
                break
            
            # Skip if we've seen a similar term
            if any(term in s or s in term for s in seen):
                continue
            
            # Generate simple definition
            definition = self._generate_definition(term, text)
            
            concepts.append(KeyConcept(
                term=term.title(),
                definition=definition,
                frequency=count,
                importance=min(1.0, count / 10)
            ))
            seen.add(term)
        
        return concepts
    
    def _generate_definition(self, term: str, text: str) -> str:
        """Generate a simple definition for a term based on context"""
        # Find sentences containing the term
        sentences = self._split_sentences(text)
        
        for sent in sentences:
            if term in sent.lower():
                # Return the sentence as context
                if len(sent) < 200:
                    return f"Context: {sent}"
        
        return "A concept mentioned in this content."
    
    def _generate_questions(
        self,
        concepts: List[KeyConcept],
        sentences: List[str],
        max_questions: int = 8
    ) -> List[StudyQuestion]:
        """Generate study questions based on content"""
        questions = []
        
        # Factual questions from concepts
        for concept in concepts[:4]:
            q = StudyQuestion(
                question=f"What is {concept.term}?",
                question_type='factual',
                difficulty='easy',
                suggested_answer=concept.definition
            )
            questions.append(q)
        
        # Conceptual questions
        if len(concepts) >= 2:
            q = StudyQuestion(
                question=f"How does {concepts[0].term} relate to {concepts[1].term}?",
                question_type='conceptual',
                difficulty='medium'
            )
            questions.append(q)
        
        # Application question
        if concepts:
            q = StudyQuestion(
                question=f"Give a real-world example of how {concepts[0].term} is used.",
                question_type='application',
                difficulty='medium'
            )
            questions.append(q)
        
        # Analysis question
        if concepts:
            q = StudyQuestion(
                question=f"What are the advantages and disadvantages of {concepts[0].term}?",
                question_type='analysis',
                difficulty='hard'
            )
            questions.append(q)
        
        # Summary question
        q = StudyQuestion(
            question="Summarize the main points discussed in this content.",
            question_type='analysis',
            difficulty='medium'
        )
        questions.append(q)
        
        return questions[:max_questions]
    
    def _find_related_topics(self, text: str, max_topics: int = 5) -> List[str]:
        """Find related topics based on content"""
        text_lower = text.lower()
        related = []
        
        # Check which domains are relevant
        for domain, keywords in self.TOPIC_KEYWORDS.items():
            matches = sum(1 for kw in keywords if kw in text_lower)
            if matches >= 2:
                related.append(domain.replace('_', ' ').title())
        
        # Add specific related topics based on content
        topic_suggestions = {
            'instagram': ['Social Media Marketing', 'Digital Photography', 'Content Creation', 'Influencer Marketing'],
            'facebook': ['Social Networking', 'Digital Advertising', 'Meta Platforms'],
            'api': ['Web Development', 'Software Integration', 'RESTful Services'],
            'database': ['Data Management', 'SQL', 'Data Architecture'],
            'machine learning': ['Artificial Intelligence', 'Data Science', 'Neural Networks'],
            'meeting': ['Project Management', 'Team Collaboration', 'Communication Skills'],
        }
        
        for keyword, topics in topic_suggestions.items():
            if keyword in text_lower:
                related.extend(topics)
        
        # Remove duplicates and limit
        seen = set()
        unique = []
        for topic in related:
            if topic.lower() not in seen:
                seen.add(topic.lower())
                unique.append(topic)
        
        return unique[:max_topics]


class DocumentGenerator:
    """
    Generates formatted documents from analyzed content
    
    Supported formats:
    - Markdown (.md)
    - HTML (.html)
    - Plain Text (.txt)
    - PDF (.pdf) - requires fpdf2
    - Word (.docx) - requires python-docx
    """
    
    def __init__(self):
        self.analyzer = ContentAnalyzer()
    
    def generate(
        self,
        text: str,
        output_path: str,
        title: str = "EchoNotes Document",
        format: str = "auto"
    ) -> str:
        """
        Generate formatted document
        
        Args:
            text: Input text/transcript
            output_path: Output file path
            title: Document title
            format: Output format (auto, md, html, txt, pdf, docx)
            
        Returns:
            Path to generated file
        """
        # Analyze content
        content = self.analyzer.analyze(text, title)
        
        # Determine format
        if format == "auto":
            ext = Path(output_path).suffix.lower()
            format = ext[1:] if ext else 'md'
        
        # Generate document
        if format == 'md':
            return self._generate_markdown(content, output_path)
        elif format == 'html':
            return self._generate_html(content, output_path)
        elif format == 'txt':
            return self._generate_text(content, output_path)
        elif format == 'pdf':
            return self._generate_pdf(content, output_path)
        elif format == 'docx':
            return self._generate_docx(content, output_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _generate_markdown(self, content: DocumentContent, path: str) -> str:
        """Generate Markdown document"""
        lines = [
            f"# {content.title}",
            "",
            f"*Generated: {content.generated_at}*",
            f"*Word Count: {content.word_count}*",
            "",
            "---",
            "",
            "## 📋 Executive Summary",
            "",
            content.executive_summary,
            "",
            "---",
            "",
            "## 🔑 Key Sentences",
            "",
        ]
        
        for i, sent in enumerate(content.key_sentences, 1):
            lines.append(f"{i}. {sent}")
        
        lines.extend([
            "",
            "---",
            "",
            "## 💡 Key Concepts",
            "",
        ])
        
        for concept in content.key_concepts:
            lines.append(f"### {concept.term}")
            lines.append(f"- **Importance**: {'⭐' * int(concept.importance * 5)}")
            lines.append(f"- {concept.definition}")
            lines.append("")
        
        lines.extend([
            "---",
            "",
            "## ❓ Study Questions",
            "",
        ])
        
        for i, q in enumerate(content.study_questions, 1):
            difficulty_emoji = {'easy': '🟢', 'medium': '🟡', 'hard': '🔴'}.get(q.difficulty, '⚪')
            lines.append(f"**{i}. {q.question}** {difficulty_emoji}")
            lines.append(f"   - Type: {q.question_type.title()}")
            if q.suggested_answer:
                lines.append(f"   - *Hint: {q.suggested_answer[:100]}...*")
            lines.append("")
        
        lines.extend([
            "---",
            "",
            "## 🔗 Related Topics",
            "",
        ])
        
        for topic in content.related_topics:
            lines.append(f"- {topic}")
        
        # Action items if present
        if content.action_items:
            lines.extend([
                "",
                "---",
                "",
                "## ✅ Action Items",
                "",
            ])
            for item in content.action_items:
                lines.append(f"- [ ] {item}")
        
        # Full content
        lines.extend([
            "",
            "---",
            "",
            "## 📄 Complete Content",
            "",
            content.full_content,
            "",
            "---",
            "",
            "*Generated by EchoNotes - Speech to Document System*",
        ])
        
        # Write file
        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return path
    
    def _generate_html(self, content: DocumentContent, path: str) -> str:
        """Generate HTML document"""
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{content.title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
            color: #333;
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        .section {{
            background: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            margin: 15px 0;
        }}
        .concept {{
            background: #e8f4fd;
            padding: 15px;
            border-left: 4px solid #3498db;
            margin: 10px 0;
        }}
        .question {{
            background: #fff3e0;
            padding: 15px;
            border-left: 4px solid #ff9800;
            margin: 10px 0;
        }}
        .difficulty {{
            display: inline-block;
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 12px;
            margin-left: 10px;
        }}
        .easy {{ background: #c8e6c9; color: #2e7d32; }}
        .medium {{ background: #fff9c4; color: #f57f17; }}
        .hard {{ background: #ffcdd2; color: #c62828; }}
        .tag {{
            display: inline-block;
            background: #e0e0e0;
            padding: 5px 12px;
            border-radius: 15px;
            margin: 5px;
            font-size: 14px;
        }}
        .metadata {{
            color: #666;
            font-size: 14px;
        }}
        ol, ul {{
            padding-left: 20px;
        }}
        li {{
            margin: 8px 0;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📝 {content.title}</h1>
        <p class="metadata">Generated: {content.generated_at} | Words: {content.word_count}</p>
        
        <h2>📋 Executive Summary</h2>
        <div class="section">
            <p>{content.executive_summary}</p>
        </div>
        
        <h2>🔑 Key Sentences</h2>
        <div class="section">
            <ol>
"""
        for sent in content.key_sentences:
            html += f"                <li>{sent}</li>\n"
        
        html += """            </ol>
        </div>
        
        <h2>💡 Key Concepts</h2>
"""
        for concept in content.key_concepts:
            stars = '⭐' * int(concept.importance * 5)
            html += f"""        <div class="concept">
            <strong>{concept.term}</strong> {stars}<br>
            <small>{concept.definition}</small>
        </div>
"""
        
        html += """        <h2>❓ Study Questions</h2>
"""
        for i, q in enumerate(content.study_questions, 1):
            html += f"""        <div class="question">
            <strong>{i}. {q.question}</strong>
            <span class="difficulty {q.difficulty}">{q.difficulty.upper()}</span><br>
            <small>Type: {q.question_type.title()}</small>
        </div>
"""
        
        html += """        <h2>🔗 Related Topics</h2>
        <div class="section">
"""
        for topic in content.related_topics:
            html += f'            <span class="tag">{topic}</span>\n'
        
        html += f"""        </div>
        
        <h2>📄 Complete Content</h2>
        <div class="section">
            <p>{content.full_content}</p>
        </div>
        
        <hr>
        <p class="metadata" style="text-align: center;">
            Generated by EchoNotes - Speech to Document System
        </p>
    </div>
</body>
</html>"""
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return path
    
    def _generate_text(self, content: DocumentContent, path: str) -> str:
        """Generate plain text document"""
        lines = [
            "=" * 60,
            f"  {content.title}",
            "=" * 60,
            f"  Generated: {content.generated_at}",
            f"  Word Count: {content.word_count}",
            "=" * 60,
            "",
            "EXECUTIVE SUMMARY",
            "-" * 40,
            content.executive_summary,
            "",
            "KEY SENTENCES",
            "-" * 40,
        ]
        
        for i, sent in enumerate(content.key_sentences, 1):
            lines.append(f"  {i}. {sent}")
        
        lines.extend([
            "",
            "KEY CONCEPTS",
            "-" * 40,
        ])
        
        for concept in content.key_concepts:
            lines.append(f"  * {concept.term}")
            lines.append(f"    {concept.definition}")
            lines.append("")
        
        lines.extend([
            "STUDY QUESTIONS",
            "-" * 40,
        ])
        
        for i, q in enumerate(content.study_questions, 1):
            lines.append(f"  {i}. [{q.difficulty.upper()}] {q.question}")
        
        lines.extend([
            "",
            "RELATED TOPICS",
            "-" * 40,
        ])
        
        for topic in content.related_topics:
            lines.append(f"  - {topic}")
        
        lines.extend([
            "",
            "COMPLETE CONTENT",
            "-" * 40,
            content.full_content,
            "",
            "=" * 60,
            "  Generated by EchoNotes",
            "=" * 60,
        ])
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return path
    
    def _generate_pdf(self, content: DocumentContent, path: str) -> str:
        """Generate PDF document"""
        try:
            from fpdf import FPDF
        except ImportError:
            print("PDF generation requires fpdf2. Install with: pip install fpdf2")
            # Fallback to text
            return self._generate_text(content, path.replace('.pdf', '.txt'))
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Title
        pdf.set_font('Helvetica', 'B', 20)
        pdf.cell(0, 15, content.title, ln=True, align='C')
        
        pdf.set_font('Helvetica', 'I', 10)
        pdf.cell(0, 8, f"Generated: {content.generated_at} | Words: {content.word_count}", ln=True, align='C')
        pdf.ln(10)
        
        # Executive Summary
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, "Executive Summary", ln=True)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 6, content.executive_summary)
        pdf.ln(5)
        
        # Key Sentences
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, "Key Sentences", ln=True)
        pdf.set_font('Helvetica', '', 11)
        for i, sent in enumerate(content.key_sentences, 1):
            pdf.multi_cell(0, 6, f"{i}. {sent}")
        pdf.ln(5)
        
        # Key Concepts
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, "Key Concepts", ln=True)
        for concept in content.key_concepts:
            pdf.set_font('Helvetica', 'B', 11)
            pdf.cell(0, 6, f"* {concept.term}", ln=True)
            pdf.set_font('Helvetica', '', 10)
            pdf.multi_cell(0, 5, f"  {concept.definition}")
        pdf.ln(5)
        
        # Study Questions
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, "Study Questions", ln=True)
        pdf.set_font('Helvetica', '', 11)
        for i, q in enumerate(content.study_questions, 1):
            pdf.multi_cell(0, 6, f"{i}. [{q.difficulty}] {q.question}")
        pdf.ln(5)
        
        # Related Topics
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, "Related Topics", ln=True)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 6, ", ".join(content.related_topics))
        
        pdf.output(path)
        return path
    
    def _generate_docx(self, content: DocumentContent, path: str) -> str:
        """Generate Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("DOCX generation requires python-docx. Install with: pip install python-docx")
            return self._generate_text(content, path.replace('.docx', '.txt'))
        
        doc = Document()
        
        # Title
        title = doc.add_heading(content.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Generated: {content.generated_at} | Words: {content.word_count}").italic = True
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(content.executive_summary)
        
        # Key Sentences
        doc.add_heading("Key Sentences", level=1)
        for i, sent in enumerate(content.key_sentences, 1):
            doc.add_paragraph(f"{i}. {sent}")
        
        # Key Concepts
        doc.add_heading("Key Concepts", level=1)
        for concept in content.key_concepts:
            p = doc.add_paragraph()
            p.add_run(f"{concept.term}: ").bold = True
            p.add_run(concept.definition)
        
        # Study Questions
        doc.add_heading("Study Questions", level=1)
        for i, q in enumerate(content.study_questions, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. [{q.difficulty.upper()}] ").bold = True
            p.add_run(q.question)
        
        # Related Topics
        doc.add_heading("Related Topics", level=1)
        doc.add_paragraph(", ".join(content.related_topics))
        
        # Full Content
        doc.add_heading("Complete Content", level=1)
        doc.add_paragraph(content.full_content)
        
        doc.save(path)
        return path
